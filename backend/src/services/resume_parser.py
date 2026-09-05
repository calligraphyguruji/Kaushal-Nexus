import io
from pathlib import Path
import re
from typing import Any, Dict, List, Optional, Tuple
import uuid

from src.core.logging import logger
from src.models.competency import Competency


class ResumeParserService:
    """
    Candidate CV text extraction, project decomposition, and skill normalization service.
    
    IMPORTANT SAFETY RULE:
    Extracted skills represent candidate self-assertions and prior evidence.
    They are NEVER written directly to Bayesian Knowledge Tracing (BKT) learner_skill_mastery.
    """

    # Robust skill synonym and alias mapping dictionary
    SKILL_ALIASES: Dict[str, Dict[str, Any]] = {
        "python": {
            "canonical": "Python Basics",
            "code_hint": "COMP-PY-BASE",
            "category": "language",
            "synonyms": ["python", "python3", "python 3", "py3", "core python"],
        },
        "python oop": {
            "canonical": "Python OOP",
            "code_hint": "COMP-PY-OOP",
            "category": "concept",
            "synonyms": ["python oop", "object oriented python", "oop concepts", "class design", "object-oriented programming"],
        },
        "sql": {
            "canonical": "SQL",
            "code_hint": "COMP-SQL-CORE",
            "category": "database",
            "synonyms": ["sql", "postgresql", "postgres", "mysql", "sqlite", "relational database", "rdbms", "queries"],
        },
        "git": {
            "canonical": "Git",
            "code_hint": "COMP-GIT-VCS",
            "category": "tool",
            "synonyms": ["git", "github", "gitlab", "version control", "vcs", "branching"],
        },
        "dsa": {
            "canonical": "DSA",
            "code_hint": "COMP-DSA-CORE",
            "category": "concept",
            "synonyms": ["dsa", "data structures", "algorithms", "problem solving", "algorithmic thinking", "leetcode"],
        },
        "rest api": {
            "canonical": "REST API",
            "code_hint": "COMP-REST-API",
            "category": "framework",
            "synonyms": ["rest api", "rest", "restful", "fastapi", "flask", "django rest framework", "drf", "api integration", "http methods", "endpoints"],
        },
        "react": {
            "canonical": "Full-Stack Web Development (React & Node)",
            "code_hint": "COMP-FULLSTACK",
            "category": "framework",
            "synonyms": ["react", "react.js", "reactjs", "next.js", "frontend development", "redux", "tailwind"],
        },
        "cloud devops": {
            "canonical": "Cloud DevOps & Kubernetes",
            "code_hint": "COMP-CLOUD-OPS",
            "category": "tool",
            "synonyms": ["docker", "kubernetes", "cloud", "aws", "azure", "gcp", "ci/cd", "devops", "containerization"],
        },
        "data analytics": {
            "canonical": "Python Application Development",
            "code_hint": "COMP-PY-DEV",
            "category": "language",
            "synonyms": ["pandas", "numpy", "matplotlib", "seaborn", "data analytics", "data cleaning", "scikit-learn"],
        },
    }

    @classmethod
    def extract_text_from_file(cls, storage_path: str, mime_type: str) -> str:
        """Extracts clean text content from PDF, DOCX, or text files."""
        path = Path(storage_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found at '{storage_path}'.")

        extracted_text = ""

        try:
            # 1. PDF extraction
            if "pdf" in mime_type.lower() or path.suffix.lower() == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                pages_text = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
                extracted_text = "\n".join(pages_text)

            # 2. DOCX extraction
            elif "word" in mime_type.lower() or path.suffix.lower() == ".docx":
                import docx
                doc = docx.Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract table text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                extracted_text = "\n".join(paragraphs)

            # 3. Plain text fallback
            else:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted_text = f.read()

        except Exception as e:
            logger.error(f"Error extracting text from file {storage_path}: {e}")
            # Try plain text fallback if PDF/Word parser choked
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted_text = f.read()
            except Exception:
                extracted_text = ""

        return extracted_text.strip()

    @classmethod
    def normalize_and_extract_skills(
        cls,
        text: str,
        available_competencies: List[Competency],
    ) -> List[Dict[str, Any]]:
        """
        Parses text for technical competencies and normalizes them against
        the standardized national competency dictionary.
        Returns a deduplicated list of candidate skills with confidence metrics.
        """
        if not text:
            return []

        lower_text = " " + text.lower() + " "
        # Build lookup maps for existing competencies
        comp_by_code: Dict[str, Competency] = {c.code.upper(): c for c in available_competencies}
        comp_by_name: Dict[str, Competency] = {c.name.lower(): c for c in available_competencies}

        extracted_skills: Dict[str, Dict[str, Any]] = {}

        # 1. Match against standardized dictionary names directly
        for comp in available_competencies:
            pattern = r"(?<!\w)" + re.escape(comp.name.lower()) + r"(?!\w)"
            if re.search(pattern, lower_text):
                extracted_skills[comp.name.lower()] = {
                    "raw_skill_text": comp.name,
                    "competency_id": comp.id,
                    "competency_code": comp.code,
                    "competency_name": comp.name,
                    "confidence": 0.95,
                    "category": "standard_competency",
                    "years_experience": None,
                }

        # 2. Match against common tech aliases and synonym clusters
        for skill_key, alias_info in cls.SKILL_ALIASES.items():
            matched_synonym = None
            for syn in alias_info["synonyms"]:
                # Use word-boundary regex
                pattern = r"(?<!\w)" + re.escape(syn.lower()) + r"(?!\w)"
                if re.search(pattern, lower_text):
                    matched_synonym = syn
                    break

            if matched_synonym:
                # Find corresponding Competency if available
                matched_comp: Optional[Competency] = None
                code_hint = alias_info.get("code_hint")
                if code_hint and code_hint.upper() in comp_by_code:
                    matched_comp = comp_by_code[code_hint.upper()]
                elif alias_info["canonical"].lower() in comp_by_name:
                    matched_comp = comp_by_name[alias_info["canonical"].lower()]

                extracted_skills[skill_key] = {
                    "raw_skill_text": matched_synonym.title(),
                    "competency_id": matched_comp.id if matched_comp else None,
                    "competency_code": matched_comp.code if matched_comp else None,
                    "competency_name": matched_comp.name if matched_comp else alias_info["canonical"],
                    "confidence": 0.90 if matched_comp else 0.80,
                    "category": alias_info.get("category", "technical"),
                    "years_experience": None,
                }

        # Estimate experience years if mentioned nearby in text
        for item in extracted_skills.values():
            skill_name = item["raw_skill_text"]
            # Look for e.g. "Python (2 years)" or "2+ years of Python"
            exp_regex = rf"(?:(\d+(?:\.\d+)?)\+?\s*years?(?:\s*of)?\s*{re.escape(skill_name.lower())}|{re.escape(skill_name.lower())}\s*\(?(\d+(?:\.\d+)?)\+?\s*years?\)?)"
            exp_match = re.search(exp_regex, lower_text)
            if exp_match:
                exp_val = exp_match.group(1) or exp_match.group(2)
                try:
                    item["years_experience"] = float(exp_val)
                except Exception:
                    pass

        return list(extracted_skills.values())

    @classmethod
    def extract_projects(cls, text: str) -> List[Dict[str, Any]]:
        """
        Extracts projects and experience entries from resume text using section headers.
        """
        if not text:
            return []

        projects: List[Dict[str, Any]] = []

        # Find project section
        section_pattern = r"(?:PROJECTS|PERSONAL PROJECTS|ACADEMIC PROJECTS|KEY PROJECTS)\s*[:\n](.*?)(?=(?:EDUCATION|EXPERIENCE|WORK EXPERIENCE|SKILLS|CERTIFICATIONS|ACHIEVEMENTS|\Z))"
        match = re.search(section_pattern, text, re.IGNORECASE | re.DOTALL)
        
        project_text = match.group(1).strip() if match else ""
        if not project_text:
            # Fallback: scan for bullet blocks or title lines
            project_text = text

        # Split project blocks (e.g. by blank lines, bullet headers, or bold lines)
        blocks = re.split(r"\n\s*\n|\n(?=[A-Z0-9][A-Za-z0-9\s\-]+(?:\s*\||\s*–|\s*\(|\s*:))", project_text)

        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            header_line = lines[0]
            # Discard trivial headers or tiny lines
            if len(header_line) < 4 or len(header_line) > 120:
                continue

            # Check if header contains a date
            date_match = re.search(r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}|\d{4})\s*(?:–|-|to)\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}|\d{4}|Present)", header_line, re.IGNORECASE)
            start_date = date_match.group(1) if date_match else None
            end_date = date_match.group(2) if date_match else None

            # Clean title
            title = re.sub(r"\(.*?\)|\|.*?$|–.*?$|-.*?$", "", header_line).strip()
            if not title:
                title = header_line[:60]

            # Description is remaining lines
            desc_lines = lines[1:]
            description = " ".join(desc_lines) if desc_lines else None

            # Detect technologies mentioned
            tech_match = re.search(r"(?:Tech(?:nologies)?|Stack|Built with|Tools?)\s*[:\-]\s*(.*?)(?:\.|\n|\Z)", block, re.IGNORECASE)
            technologies = tech_match.group(1).strip() if tech_match else None

            projects.append({
                "title": title[:200],
                "description": description[:1000] if description else None,
                "technologies": technologies[:500] if technologies else None,
                "start_date": start_date,
                "end_date": end_date,
            })

            if len(projects) >= 6:
                break

        return projects


resume_parser = ResumeParserService()
